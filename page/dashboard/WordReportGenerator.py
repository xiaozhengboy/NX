import io
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt, Cm, Mm, Inches
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path
from page.dashboard.LoggingDashboard import logger
from page.dashboard import Utile
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


class WordReportGenerator:
    """Word报告生成器（按照alarm_analyzer.py样式优化）"""

    def __init__(self, alert_dir='alerts'):
        self.alert_dir = Path(alert_dir)
        self.logo_path = Path("static") / "logo.png"

    def set_font_style(self, run, font_name="SimHei", size=Pt(10), bold=False, color=None):
        """统一设置字体样式的工具函数"""
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if size:
            run.font.size = size
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def set_table_unified_style(self, table, header_row_count=1, header_font_size=Pt(10), content_font_size=Pt(9)):
        """
        统一设置表格样式（参考alarm_analyzer.py）
        :param table: Word表格对象
        :param header_row_count: 表头行数（默认1行）
        :param header_font_size: 表头字体大小
        :param content_font_size: 内容行字体大小
        """
        # 表格整体对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False

        # 遍历所有行，统一设置样式
        for row_idx, row in enumerate(table.rows):
            # 设置行高
            row.height = Mm(8)  # 统一行高8毫米
            row.height_rule = 1  # 固定行高

            # 遍历当前行的所有单元格
            for cell in row.cells:
                # 单元格垂直居中
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                # 遍历单元格内的段落
                for para in cell.paragraphs:
                    # 水平居中
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # 统一设置字体样式
                    for run in para.runs:
                        if row_idx < header_row_count:
                            # 表头样式：加粗、指定字号
                            self.set_font_style(run, size=header_font_size, bold=True)
                        else:
                            # 内容行样式：常规、指定字号
                            self.set_font_style(run, size=content_font_size, bold=False)

    def generate_word_report(self, alerts, search_params, include_data=False, include_images=False):
        """
        生成Word格式的报告
        :param alerts: 告警数据列表
        :param search_params: 搜索参数
        :param include_data: 是否包含源数据Excel
        :param include_images: 是否包含图片
        :return: (word_buffer, excel_buffer, image_infos) 元组
        """
        try:
            # 1. 生成Word文档
            word_buffer = self._create_word_document(alerts, search_params, include_images)

            # 2. 如果需要，生成Excel数据
            excel_buffer = None
            if include_data:
                excel_buffer = self._create_excel_data(alerts, search_params)

            # 3. 准备图片信息
            image_infos = []
            if include_images:
                image_infos = self._prepare_image_infos(alerts)

            return word_buffer, excel_buffer, image_infos

        except Exception as e:
            logger.error(f"生成报告失败: {e}")
            raise e

    def _create_word_document(self, alerts, search_params, include_images):
        """创建Word文档（按照alarm_analyzer.py样式）"""
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENTATION, WD_SECTION
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.shared import Inches, RGBColor, Pt, Cm
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("缺少python-docx库，请安装: pip install python-docx")
            raise

        # 创建Word文档
        doc = Document()

        # ========== 封页设置（第一节） ==========
        first_section = doc.sections[0]
        first_section.different_first_page_header_footer = True
        first_section.header.is_linked_to_previous = False
        first_section.footer.is_linked_to_previous = False
        first_section.page_numbering_start = 0
        first_section.restart_page_numbering = True

        # 设置封页页面样式
        first_section.orientation = WD_ORIENTATION.PORTRAIT
        first_section.left_margin = Cm(2.5)
        first_section.right_margin = Cm(2.5)
        first_section.top_margin = Cm(3)
        first_section.bottom_margin = Cm(2.5)

        # 1. 左上角插入logo
        if self.logo_path.exists():
            logo_para = doc.add_paragraph()
            logo_run = logo_para.add_run()
            logo_run.add_picture(str(self.logo_path), width=Cm(6), height=Cm(2))
            logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for _ in range(6):
                doc.add_paragraph()
        else:
            logger.warning(f"未找到logo文件：{self.logo_path}，跳过logo插入")
            for _ in range(8):
                doc.add_paragraph()

        # 2. 封页标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run('叶片检测报告')


        for run in title_para.runs:
            self.set_font_style(run, size=Pt(20), bold=True)

        # 3. 巡检时间
        time_para = doc.add_paragraph()
        # 格式化日期
        start_time = search_params.get('start_time', '')
        end_time = search_params.get('end_time', '')

        # 转换日期格式
        try:
            start_dt = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
            end_dt = datetime.fromisoformat(end_time.replace('Z', '+00:00'))
            time_text = f"{start_dt.strftime('%Y-%m-%d %H:%M')} - {end_dt.strftime('%Y-%m-%d %H:%M')}"
        except Exception as e:
            time_text = f"{start_time} - {end_time}"

        time_run = time_para.add_run(time_text)
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.set_font_style(time_run, size=Pt(14), color=RGBColor(64, 64, 64))

        # 4. 封页底部空行 + 分节符
        for _ in range(8):
            doc.add_paragraph()
        doc.add_section(WD_SECTION.NEW_PAGE)

        # ========== 正文设置（第二节及以后） ==========
        main_section = doc.sections[1]
        main_section.odd_and_even_pages_header_footer = True
        main_section.different_first_page_header_footer = False
        main_section.header.is_linked_to_previous = False
        main_section.footer.is_linked_to_previous = False
        main_section.page_numbering_start = 1
        main_section.restart_page_numbering = True

        # 正文页面样式
        main_section.orientation = WD_ORIENTATION.PORTRAIT
        main_section.left_margin = Cm(2.5)
        main_section.right_margin = Cm(2.5)
        main_section.top_margin = Cm(3)
        main_section.bottom_margin = Cm(2.5)

        # ========== 正文页眉设置 ==========
        # 奇数页页眉
        odd_header = main_section.header
        odd_header_para = odd_header.add_paragraph('叶片检测报告')
        odd_header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in odd_header_para.runs:
            self.set_font_style(run, size=Pt(10))

        # 偶数页页眉
        even_header = main_section.even_page_header
        even_header_para = even_header.add_paragraph('叶片检测报告')
        even_header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in even_header_para.runs:
            self.set_font_style(run, size=Pt(10))

        # ========== 正文内容 ==========
        # 1. 巡检综述
        doc.add_paragraph()
        review_heading = doc.add_heading('1 巡检综述', level=1)
        for run in review_heading.runs:
            self.set_font_style(run, size=Pt(14), bold=True)

        # ------------ 定义固定宽度 ------------
        TOTAL_WIDTH = 160  # 总宽度160毫米（页面宽度减去边距）
        # 2列表格：1:4宽度（32mm + 128mm）
        COL1_WIDTH = Mm(32)
        COL2_WIDTH = Mm(128)
        # 5列表格：平均分（32mm/列 ×5）
        AVG_COL_WIDTH = Mm(32)

        # ------------ 2列表格 ------------
        overview_table1 = doc.add_table(rows=5, cols=2, style='Table Grid')
        overview_table1.autofit = False
        overview_table1.allow_autofit = False
        overview_table1.preferred_width = Mm(TOTAL_WIDTH)
        overview_table1.preferred_width_type = 2

        # 填充内容
        overview_data1 = [
            ('站点名称', '宁夏宁东五六七八期升压站'),
            ('任务名称', '全站叶片检测'),
            ('任务类型', '自动巡视'),
            ('开始时间', start_time),
            ('结束时间', end_time)
        ]

        for row_idx, (label, value) in enumerate(overview_data1):
            cell1 = overview_table1.cell(row_idx, 0)
            cell2 = overview_table1.cell(row_idx, 1)

            # 清空原有内容
            cell1.text = ''
            cell2.text = ''

            # 添加内容
            p1 = cell1.paragraphs[0]
            p2 = cell2.paragraphs[0]
            p1.add_run(label)
            p2.add_run(str(value))

            # 设置单元格宽度
            cell1.width = COL1_WIDTH
            cell2.width = COL2_WIDTH
            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 应用样式
        self.set_table_unified_style(overview_table1, header_row_count=0, content_font_size=Pt(9))
        overview_table1.alignment = WD_TABLE_ALIGNMENT.CENTER


        # ------------ 5列表格 ------------
        # 计算统计信息
        stats = self._calculate_statistics(alerts)
        overview_table2 = doc.add_table(rows=2, cols=5, style='Table Grid')
        overview_table2.autofit = False
        overview_table2.allow_autofit = False
        overview_table2.preferred_width = Mm(TOTAL_WIDTH)
        overview_table2.preferred_width_type = 2

        # 填充内容
        overview_data2 = [
            ('巡检设备数', '缺陷设备数', '总巡检图像', '有缺陷图像', '总缺陷数量'),
            (
                str(stats.get('涉及风机数', 0)),
                str(stats.get('缺陷设备数', 0)),
                str(stats.get('总巡检图像', 0)),
                str(stats.get('有缺陷图像', 0)),
                str(stats.get('总缺陷数量', 0))
            )
        ]

        for row_idx, row_data in enumerate(overview_data2):
            for col_idx, value in enumerate(row_data):
                cell = overview_table2.cell(row_idx, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                p.add_run(str(value))
                cell.width = AVG_COL_WIDTH
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 应用样式
        self.set_table_unified_style(overview_table2, header_row_count=1, content_font_size=Pt(9))
        overview_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 2. 设备维度统计表格
        doc.add_paragraph()
        device_heading = doc.add_heading('2 风机缺陷类型汇总', level=1)
        for run in device_heading.runs:
            self.set_font_style(run, size=Pt(14), bold=True)

        # 收集所有风机和缺陷类型
        cameras = set()
        defect_types = set()
        camera_defect_counts = {}

        for alert in alerts:
            camera_id = alert.get('camera_id')
            if camera_id:
                cameras.add(camera_id)
                if camera_id not in camera_defect_counts:
                    camera_defect_counts[camera_id] = {}

                for detection in alert.get('detections', []):
                    defect_name = detection.get('name')
                    if defect_name:
                        defect_types.add(defect_name)
                        camera_defect_counts[camera_id][defect_name] = \
                            camera_defect_counts[camera_id].get(defect_name, 0) + 1

        # 排序
        sorted_cameras = sorted(list(cameras))
        sorted_defects = sorted(list(defect_types), key=lambda x: Utile.translate_defect_name(x))

        # 创建表格
        device_table = doc.add_table(rows=len(sorted_cameras) + 1, cols=len(sorted_defects) + 2)
        device_table.style = 'Table Grid'
        device_table.autofit = False
        device_table.allow_autofit = False
        device_table.preferred_width = Mm(TOTAL_WIDTH)
        device_table.preferred_width_type = 2

        # 设置表头
        header_cells = device_table.rows[0].cells
        header_cells[0].text = "设施名称"
        header_cells[1].text = "总数"

        # 设置表头列宽
        name_width = Mm(40)  # 风机名称列宽
        total_width = Mm(20)  # 总数列宽
        defect_width = Mm((TOTAL_WIDTH - 60) / len(sorted_defects)) if sorted_defects else Mm(20)

        header_cells[0].width = name_width
        header_cells[1].width = total_width

        for i, defect in enumerate(sorted_defects, 2):
            header_cells[i].text = Utile.translate_defect_name(defect)
            header_cells[i].width = defect_width

        # 填充数据
        for row_idx, camera in enumerate(sorted_cameras, 1):
            row_cells = device_table.rows[row_idx].cells
            row_cells[0].text = camera
            row_cells[0].width = name_width

            # 计算该风机的总缺陷数
            total_defects = sum(camera_defect_counts.get(camera, {}).values())
            row_cells[1].text = str(total_defects) if total_defects > 0 else "0"
            row_cells[1].width = total_width

            # 填充各缺陷类型数量
            for col_idx, defect in enumerate(sorted_defects, 2):
                count = camera_defect_counts.get(camera, {}).get(defect, 0)
                row_cells[col_idx].text = str(count) if count > 0 else "0"
                row_cells[col_idx].width = defect_width

        # 应用样式
        self.set_table_unified_style(device_table, header_row_count=1, content_font_size=Pt(9))

        # 3. 缺陷类型汇总表格
        doc.add_paragraph()
        alarm_heading = doc.add_heading('3 各缺陷类型风机汇总', level=1)
        for run in alarm_heading.runs:
            self.set_font_style(run, size=Pt(14), bold=True)

        # 统计缺陷类型数据
        defect_camera_counts = {}

        for alert in alerts:
            camera_id = alert.get('camera_id')
            for detection in alert.get('detections', []):
                defect_name = detection.get('name')
                if defect_name:
                    if defect_name not in defect_camera_counts:
                        defect_camera_counts[defect_name] = {
                            'cameras': set(),
                            'count': 0
                        }

                    defect_camera_counts[defect_name]['cameras'].add(camera_id)
                    defect_camera_counts[defect_name]['count'] += 1

        # 排序缺陷类型
        sorted_defects = sorted(defect_camera_counts.items(),
                                key=lambda x: Utile.translate_defect_name(x[0]))

        # 创建表格
        alarm_table = doc.add_table(rows=len(sorted_defects) + 1, cols=3)
        alarm_table.style = 'Table Grid'
        alarm_table.autofit = False
        alarm_table.allow_autofit = False
        alarm_table.preferred_width = Mm(TOTAL_WIDTH)
        alarm_table.preferred_width_type = 2

        # 设置列宽
        defect_width = Mm(50)
        distribution_width = Mm(80)
        count_width = Mm(30)

        # 设置表头
        alarm_hdr_cells = alarm_table.rows[0].cells
        alarm_hdr_cells[0].text = "缺陷类型"
        alarm_hdr_cells[1].text = "分布"
        alarm_hdr_cells[2].text = "数量"

        alarm_hdr_cells[0].width = defect_width
        alarm_hdr_cells[1].width = distribution_width
        alarm_hdr_cells[2].width = count_width

        # 填充数据
        for row_idx, (defect_name, data) in enumerate(sorted_defects, 1):
            row_cells = alarm_table.rows[row_idx].cells

            # 缺陷类型（中文名）
            row_cells[0].text = Utile.translate_defect_name(defect_name)
            row_cells[0].width = defect_width

            # 分布（风机列表）
            cameras = sorted(list(data['cameras']))
            row_cells[1].text = '、'.join(cameras) if cameras else ""
            row_cells[1].width = distribution_width

            # 数量
            row_cells[2].text = str(data['count'])
            row_cells[2].width = count_width

        # 应用样式
        self.set_table_unified_style(alarm_table, header_row_count=1, content_font_size=Pt(9))

        # 4. 巡检总结
        doc.add_paragraph()
        summary_heading = doc.add_heading('4 巡检总结', level=1)
        for run in summary_heading.runs:
            self.set_font_style(run, size=Pt(14), bold=True)

        # 计算告警次数最多的设备
        device_total_counts = {}
        for camera_id, defect_counts in camera_defect_counts.items():
            device_total_counts[camera_id] = sum(defect_counts.values())

        top_device = None
        top_device_count = 0
        if device_total_counts:
            top_device = max(device_total_counts, key=device_total_counts.get)
            top_device_count = device_total_counts[top_device]

        # 构建总结文本
        total_devices = len(camera_defect_counts)
        total_alarm_types = len(defect_types)
        total_alarm_times = sum(device_total_counts.values())

        # 计算各告警类型的总次数
        alarm_type_counts = {}
        for defect_name in defect_types:
            total = sum(
                camera_defect_counts.get(camera_id, {}).get(defect_name, 0)
                for camera_id in camera_defect_counts
            )
            alarm_type_counts[defect_name] = total

        # 排序告警类型（按次数降序）
        sorted_alarm_types = sorted(alarm_type_counts.items(), key=lambda x: x[1], reverse=True)

        # 构建告警类型描述文本
        alarm_type_desc = []
        for i, (alarm_type, count) in enumerate(sorted_alarm_types):
            if i < 3:  # 重点显示前3种
                alarm_type_desc.append(f"{Utile.translate_defect_name(alarm_type)}产生{count}次")
            elif i == 3 and len(sorted_alarm_types) > 4:
                remaining = len(sorted_alarm_types) - 3
                alarm_type_desc.append(f"其他{remaining}种缺陷共产生{sum(c for t, c in sorted_alarm_types[3:])}次")
                break
            elif i == len(sorted_alarm_types) - 1 and len(sorted_alarm_types) <= 4:
                alarm_type_desc.append(f"{Utile.translate_defect_name(alarm_type)}产生{count}次")

        # 添加总结段落
        stat_para = doc.add_paragraph()
        stat_text = f"本次巡检任务中，共有 {total_devices} 台设备产生告警，累计产生告警 {total_alarm_times} 次，涉及 {total_alarm_types} 种告警类型，"
        if alarm_type_desc:
            stat_text += f"其中{'; '.join(alarm_type_desc)}。"
        else:
            stat_text += "。"

        if top_device:
            stat_text += f"从设备维度看，{top_device}产生的告警次数最多，共计{top_device_count}次。"

        stat_run = stat_para.add_run(stat_text)
        self.set_font_style(stat_run, size=Pt(10))

        # 分页
        doc.add_page_break()

        # 5. 附图详情
        doc.add_paragraph()
        image_heading = doc.add_heading('5 附图详情', level=1)
        for run in image_heading.runs:
            self.set_font_style(run, size=Pt(14), bold=True)

        # 按风机号分组
        camera_alerts = {}
        for alert in alerts:
            camera_id = alert.get('camera_id')
            if camera_id:
                if camera_id not in camera_alerts:
                    camera_alerts[camera_id] = []
                camera_alerts[camera_id].append(alert)

        # 按风机号排序
        sorted_cameras = sorted(camera_alerts.keys())

        for camera_id in sorted_cameras:
            # 添加风机号标题
            camera_heading = doc.add_heading(f'{camera_id}', level=2)
            for run in camera_heading.runs:
                self.set_font_style(run, size=Pt(12), bold=True)

            local_img_infos = []
            # 修复：使用正确的方式获取图片路径
            for alert in camera_alerts[camera_id]:
                # 尝试多个可能的字段名获取图片路径
                image_paths = [
                    alert.get('image_filename'),
                    alert.get('relative_path', '') + '/images/' + alert.get('alert_id', '') + '.jpg'
                ]

                found = False
                for img_path in image_paths:
                    if img_path:
                        if 'images/' in img_path:
                            # 直接使用完整路径
                            image_full_path = self.alert_dir / img_path
                        else:
                            # 否则尝试从relative_path构建
                            relative_path = alert.get('relative_path', '')
                            if relative_path:
                                image_full_path = self.alert_dir / relative_path / 'images' / f"{alert.get('alert_id', '')}.jpg"
                            else:
                                # 最后尝试使用默认路径
                                image_full_path = self.alert_dir / camera_id / 'images' / f"{alert.get('alert_id', '')}.jpg"

                        if image_full_path.exists():
                            found = True
                            # 获取缺陷名称中文
                            defect_names = []
                            for detection in alert.get('detections', []):
                                name_chinese = detection.get('name_chinese')
                                if not name_chinese:
                                    name_chinese = Utile.translate_defect_name(detection.get('name', ''))
                                if name_chinese:
                                    defect_names.append(name_chinese)

                            local_img_infos.append({
                                'local_path': image_full_path,
                                'active_name': '、'.join(defect_names) if defect_names else '未知缺陷',
                                'event_time': alert.get('detection_time', ''),
                                'security_name': alert.get('alert_level', '中等') or '中等'  # 使用告警等级
                            })
                            break

                if not found:
                    logger.warning(f"未找到设备[{camera_id}]告警[{alert.get('alert_id')}]的图片")

            if not local_img_infos and include_images:
                no_image_para = doc.add_paragraph(f"{camera_id}暂无可用缺陷图片")
                for run in no_image_para.runs:
                    self.set_font_style(run, size=Pt(9), color=RGBColor(128, 128, 128))
                continue

            # 设备内图片按事件时间升序排列
            sorted_imgs = sorted(local_img_infos, key=lambda x: x["event_time"])

            if include_images:
                for idx, img_info in enumerate(sorted_imgs, 1):
                    try:
                        # 添加图片
                        img_path = str(img_info['local_path'])
                        if not Path(img_path).exists():
                            logger.error(f"图片文件不存在: {img_path}")
                            continue

                        pic = doc.add_picture(img_path, width=Inches(6))
                        width, height = pic.width, pic.height
                        max_width = Inches(6)
                        if width > max_width:
                            ratio = max_width / width
                            pic.width = max_width
                            pic.height = int(height * ratio)

                        # 图片描述
                        img_desc = (
                            f"缺陷名称：{img_info['active_name']} | "
                            f"检测时间：{img_info['event_time']} | "
                            f"缺陷程度：{img_info['security_name']}"
                        )
                        img_para = doc.add_paragraph()
                        img_run = img_para.add_run(img_desc)
                        img_run.bold = True
                        self.set_font_style(img_run, size=Pt(9))

                    except Exception as e:
                        logger.error(f"Word中插入设备[{camera_id}]图片{idx}失败：{str(e)}")
                        continue
            else:
                # 如果不包含图片，显示文字说明
                for idx, img_info in enumerate(sorted_imgs, 1):
                    img_desc = (
                        f"图片{idx}: 缺陷名称：{img_info['active_name']} | "
                        f"检测时间：{img_info['event_time']} | "
                        f"缺陷程度：{img_info['security_name']}"
                    )
                    img_para = doc.add_paragraph()
                    img_run = img_para.add_run(img_desc)
                    self.set_font_style(img_run, size=Pt(9), color=RGBColor(128, 128, 128))

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer



    def _calculate_statistics(self, alerts):
        """计算统计信息"""
        stats = {
            '涉及风机数': 0,
            '缺陷设备数': 0,
            '总巡检图像': len(alerts),
            '有缺陷图像': 0,
            '总缺陷数量': 0
        }

        if not alerts:
            return stats

        # 计算涉及风机数
        cameras = set()
        defective_cameras = set()
        total_defects = 0
        images_with_defects = 0

        for alert in alerts:
            camera_id = alert.get('camera_id')
            if camera_id:
                cameras.add(camera_id)

            detections = alert.get('detections', [])
            if detections:
                images_with_defects += 1
                if camera_id:
                    defective_cameras.add(camera_id)
                total_defects += len(detections)

        stats['涉及风机数'] = len(cameras)
        stats['缺陷设备数'] = len(defective_cameras)
        stats['有缺陷图像'] = images_with_defects
        stats['总缺陷数量'] = total_defects

        return stats

    def _create_excel_data(self, alerts, search_params):
        """创建Excel数据文件"""
        if not alerts:
            return None

        try:
            # 准备数据
            data = []
            for alert in alerts:
                for det in alert.get('detections', []):
                    row = {
                        '告警ID': alert.get('alert_id', ''),
                        '风机编号': alert.get('camera_id', ''),
                        '检测时间': alert.get('detection_time', ''),
                        '缺陷名称': det.get('name_chinese', det.get('name', '')),
                        '缺陷英文名': det.get('name', ''),
                        '置信度': det.get('conf', 0),
                        'X坐标': det.get('x', 0),
                        'Y坐标': det.get('y', 0),
                        '宽度': det.get('w', 0),
                        '高度': det.get('h', 0),
                        '旋转角度': det.get('r', 0),
                        '图片路径': alert.get('image_filename', '')
                    }
                    data.append(row)

            # 创建DataFrame
            df = pd.DataFrame(data)

            # 保存到内存
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='告警数据', index=False)

                # 添加搜索条件工作表
                condition_data = {
                    '搜索条件': ['开始时间', '结束时间', '风机编号', '缺陷类型', '最低置信度'],
                    '值': [
                        search_params.get('start_time', ''),
                        search_params.get('end_time', ''),
                        search_params.get('camera_id', '全部'),
                        search_params.get('defect_name', '全部'),
                        search_params.get('min_confidence', 0)
                    ]
                }
                condition_df = pd.DataFrame(condition_data)
                condition_df.to_excel(writer, sheet_name='搜索条件', index=False)

            buffer.seek(0)
            return buffer

        except Exception as e:
            logger.error(f"创建Excel数据失败: {e}")
            return None

    def _prepare_image_infos(self, alerts):
        """准备图片信息"""
        image_infos = []

        for alert in alerts:
            camera_id = alert.get('camera_id', '未知')
            alert_id = alert.get('alert_id', '未知')

            # 修复：多种方式尝试获取图片路径
            image_paths = [
                alert.get('image_filename'),
                alert.get('image_path'),
                alert.get('relative_path', '') + '/images/' + alert.get('alert_id', '') + '.jpg'
            ]

            for img_path in image_paths:
                if img_path:
                    try:
                        source_path = self.alert_dir / img_path
                        if not source_path.exists():
                            # 尝试从relative_path构建路径
                            relative_path = alert.get('relative_path', '')
                            if relative_path:
                                source_path = self.alert_dir / relative_path / 'images' / f"{alert_id}.jpg"

                        if source_path.exists():
                            # 创建有意义的文件名
                            timestamp = alert.get('detection_time', '').replace(':', '').replace(' ', '_')

                            # 确保文件名安全
                            safe_camera_id = camera_id.replace('/', '_').replace('\\', '_')
                            safe_timestamp = timestamp.replace('/', '_').replace('\\', '_')

                            ext = source_path.suffix or '.jpg'
                            dest_name = f"{safe_camera_id}_{safe_timestamp}_{alert_id}{ext}"

                            image_infos.append({
                                'source_path': source_path,
                                'filename_in_zip': f"缺陷图片/{dest_name}",
                                'camera_id': camera_id,
                                'alert_id': alert_id,
                                'timestamp': timestamp
                            })
                            break  # 找到图片后跳出循环

                    except Exception as e:
                        logger.warning(f"准备图片信息失败: {e}")
                        continue

        return image_infos